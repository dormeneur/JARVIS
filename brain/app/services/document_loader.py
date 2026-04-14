"""Document loader service for extracting text from vault files."""

import csv
import hashlib
import io
import json
import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Generator, Optional

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)

# Configuration constants
SUPPORTED_EXTENSIONS = [".md", ".txt", ".pdf", ".json", ".csv", ".docx"]
MAX_FILE_SIZE_MB = 1
EXCLUDED_FOLDERS = ["Secrets", "system", ".system", ".git", "__pycache__", "node_modules"]
EXCLUDED_FILES = [".gitkeep", "file_versions.db", "devices.json"]


@dataclass
class LoadedDocument:
    """Represents a loaded document with metadata."""
    path: str
    content: str
    last_modified: datetime
    content_hash: str


class DocumentLoader:
    """Loads and extracts text from vault files."""
    
    def __init__(self, vault_path: str):
        """Initialize document loader.
        
        Args:
            vault_path: Path to the vault directory (e.g., /JARVIS)
        """
        self.vault_path = Path(vault_path)
    
    def load_documents(self) -> Generator[LoadedDocument, None, None]:
        """Walk vault and yield LoadedDocument instances.
        
        Yields:
            LoadedDocument instances for each supported file
        """
        for file_path in self._walk_vault():
            if self._should_skip(file_path):
                continue
            
            try:
                content = self._extract_content(file_path)
                if content is None:
                    continue
                
                yield LoadedDocument(
                    path=str(file_path.relative_to(self.vault_path)),
                    content=content,
                    last_modified=datetime.fromtimestamp(file_path.stat().st_mtime),
                    content_hash=self._compute_hash(file_path)
                )
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")
                continue
    
    def _walk_vault(self) -> Generator[Path, None, None]:
        """Walk vault directory recursively, pruning excluded folders.
        
        Yields:
            Path objects for each file
        """
        import os
        if not self.vault_path.exists():
            logger.warning(f"Vault path does not exist: {self.vault_path}")
            return
        
        vault_root = str(self.vault_path)
        for root, dirs, files in os.walk(vault_root):
            # Prune excluded directories in-place to prevent os.walk from entering them
            dirs[:] = [d for d in dirs if d not in EXCLUDED_FOLDERS]
            
            for file in files:
                yield Path(root) / file
    
    def _should_skip(self, path: Path) -> bool:
        """Check if file should be excluded.
        
        Args:
            path: File path to check
            
        Returns:
            True if file should be skipped, False otherwise
        """
        # Check folder exclusions
        for part in path.parts:
            if part in EXCLUDED_FOLDERS:
                return True
        
        # Check filename exclusions
        if path.name in EXCLUDED_FILES:
            return True
        
        # Check file size
        try:
            if path.stat().st_size > MAX_FILE_SIZE_MB * 1024 * 1024:
                logger.info(f"Skipping large file: {path} ({path.stat().st_size / (1024*1024):.1f} MB)")
                return True
        except OSError as e:
            logger.error(f"Failed to stat {path}: {e}")
            return True
        
        # Check extension — allow extensionless files (treat as text)
        if path.suffix and path.suffix not in SUPPORTED_EXTENSIONS:
            return True
        
        return False
    
    def _extract_content(self, file_path: Path) -> Optional[str]:
        """Extract text content from file based on type.
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text content or None if extraction failed
        """
        extension = file_path.suffix.lower()
        
        try:
            if extension in [".md", ".txt", ""]:
                return self._extract_text(file_path)
            elif extension == ".pdf":
                return self._extract_pdf(file_path)
            elif extension == ".docx":
                return self._extract_docx(file_path)
            elif extension == ".json":
                return self._extract_json(file_path)
            elif extension == ".csv":
                return self._extract_csv(file_path)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return None
        except Exception as e:
            logger.error(f"Failed to extract content from {file_path}: {e}")
            return None
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from markdown or txt file.
        
        Args:
            file_path: Path to file
            
        Returns:
            File content as UTF-8 text
        """
        return file_path.read_text(encoding="utf-8")
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file using PyMuPDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        doc = fitz.open(file_path)
        text_parts = []
        
        for page in doc:
            text_parts.append(page.get_text())
        
        doc.close()
        return "\n".join(text_parts)
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    
    def _extract_json(self, file_path: Path) -> Optional[str]:
        """Parse and format JSON file.
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            Pretty-printed JSON or None if invalid
        """
        try:
            content = file_path.read_text(encoding="utf-8")
            data = json.loads(content)
            return json.dumps(data, indent=2)
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in {file_path}: {e}")
            return None
    
    def _extract_csv(self, file_path: Path) -> Optional[str]:
        """Parse CSV and convert to markdown table.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Markdown table or None if invalid
        """
        try:
            content = file_path.read_text(encoding="utf-8")
            reader = csv.reader(io.StringIO(content))
            rows = list(reader)
            
            if not rows:
                return ""
            
            # Header row
            markdown = "| " + " | ".join(rows[0]) + " |\n"
            # Separator row
            markdown += "| " + " | ".join(["---"] * len(rows[0])) + " |\n"
            # Data rows
            for row in rows[1:]:
                # Pad row if it has fewer columns than header
                padded_row = row + [""] * (len(rows[0]) - len(row))
                markdown += "| " + " | ".join(padded_row[:len(rows[0])]) + " |\n"
            
            return markdown
        except Exception as e:
            logger.error(f"Invalid CSV in {file_path}: {e}")
            return None
    
    def _compute_hash(self, file_path: Path) -> str:
        """Compute SHA-256 hash of raw file bytes.
        
        Args:
            file_path: Path to file
            
        Returns:
            Hexadecimal hash string
        """
        sha256 = hashlib.sha256()
        
        with open(file_path, "rb") as f:
            # Read in chunks to handle large files
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        
        return sha256.hexdigest()
